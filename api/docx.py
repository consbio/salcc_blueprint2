import re

from docx import Document


# Example placeholder: {{label:title}}
PLACEHOLDER_REGEX = re.compile(r'{{(?P<scope>\w+):(?P<key>\w+)}}')


def create_docx_report(path, context):
    """
        Creates reports for analyses in docx format.
    """
    doc = Document('template.docx')

    for p in doc.paragraphs:
        # Assumption: placeholders are always completely contained within a run
        for r in p.runs:
            # a run may have multiple placeholders
            for match in re.finditer(PLACEHOLDER_REGEX, r.text):
                scope, key = match.groups()

                item = _resolve(scope, key, context)
                if isinstance(item, str):
                    # match.group() has the original text to replace
                    r.text = r.text.replace(match.group(), item)

                elif scope == 'image':
                    # Clear out the placeholder
                    r.text = r.text.replace(match.group(), '')
                    r.add_picture()

                elif scope == 'table':
                    r.text = ""
                    # do something

        if '{{ECOSYSTEMS}}' in p.text:
            # Remove '{{ECOSYSTEMS}}'placeholder and add the report content above empty p.text
            p.text = ""
            # do things
        elif '{{PARTNERS}}' in p.text:
            # Remove '{{PARTNERS}}'placeholder and add the report content above empty p.text
            p.text = ""
            # do things

    doc.save(path)


def _resolve(scope, key, context):
    """
    Resolve scope and key to a context item
    Provides very minimal validation of presence

    Parameters
    ----------
    scope: str
        singular variant of scope in context
    key: str
        key to lookup context item in context within scope
    context: dict

    Returns
    -------
    context item: str, dict, etc
    """

    if scope not in context:
        raise ValueError('Scope {0} is not found in context!'.format(scope))

    if key not in context[scope]:
        raise ValueError('Key {0} is not found in context!'.format(key))

    return context[scope][key]