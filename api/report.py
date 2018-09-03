import re
import os
import json
import docx
import matplotlib

matplotlib.use("Agg")

from collections import defaultdict
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from api.map_client import *
from api.charts import get_pie_chart, get_line_chart

DATA_DIR = os.getenv("DATA_DIR", "./public/data")

TEMPLATE = "api/template.docx"

# TODO: Improve custom table style in template.docx
DOC_STYLE = "SALCC_style"

# Example placeholder: {{label:title}}
PLACEHOLDER_REGEX = re.compile(r"{{(?P<scope>\w+):(?P<key>\w+)}}")

TOTAL_ROW_COLOR = "cacdd1"


def create_report(unit_id, path, config):
    """
        Creates report for summary unit in docx format.

    Parameters
    ----------
    unit_id: string id of a summary unit
    path: string location where docx to be saved
    config: dict of json data from src/config

    Returns
    -------
    report: docx file

    """
    doc = Document(TEMPLATE)

    context = generate_report_context(unit_id, config)
    table_counter = 0

    for p in doc.paragraphs:
        # Assumption: placeholders are always completely contained within a run
        for r in p.runs:
            # a run may have multiple placeholders
            for match in re.finditer(PLACEHOLDER_REGEX, r.text):
                scope, key = match.groups()

                item = _resolve(scope, key, context)

                if isinstance(item, str):
                    # match.group() has the original text to replace
                    if scope == "table":
                        r.add_break(WD_BREAK.LINE)
                    r.text = r.text.replace(match.group(), item)

                    if scope in {"table", "chart"}:
                        r.font.italic = True

                elif scope in {"map", "chart"}:
                    # Pictures are added at the run level
                    r.text = ""
                    size = 11.0
                    if scope == "map":
                        size = 16.5
                    if scope == "chart" and key == "priorities":  # If chart is present, put on its own page
                        r.add_break(WD_BREAK.PAGE)
                    if not isinstance(context[scope][key], str):
                        run = p.add_run()
                        run.add_picture(context[scope][key], width=Cm(size))

                        # add a caption if we have one in context
                        caption_name = "{0}_{1}".format(scope, key)
                        # if key in context["caption"]:
                        if caption_name in context["caption"]:
                            caption = doc.add_paragraph(
                                context["caption"][caption_name], style="TableCaption"
                            )
                            _move_p_after(caption, p)

                elif scope == "table":
                    r.text = ""
                    table_counter += 1
                    create_table(doc, item, p)

        if "{{INDICATORS}}" in p.text:
            # Remove '{{INDICATORS}}'placeholder and add the report content above empty p.text
            p.text = ""

            for ecosystem in context["ecosystem_indicators"]:

                if ecosystem["indicators"]:

                    name = ecosystem["ecosystem_name"]
                    percent = ecosystem["ecosystem_percentage"]
                    if percent is not "":
                        heading = "{0}: {1}% of area".format(name, percent)
                        p.insert_paragraph_before(heading, style="Heading13")
                    else:
                        heading = name
                        p.insert_paragraph_before(heading, style="Heading13")

                    for indicator in ecosystem["indicators"]:
                        table_counter += 1

                        p.insert_paragraph_before(
                            indicator["value"]["indicator_name"], style="Heading11"
                        )
                        p.insert_paragraph_before(
                            indicator["value"]["indicator_description"]
                        )

                        caption_text = "Table {0}: {1}".format(
                            table_counter, indicator["value"]["indicator_caption"]
                        )
                        caption = p.insert_paragraph_before(
                            caption_text, style="TableCaption"
                        )

                        create_table(
                            doc,
                            indicator["table"]["indicator_table"],
                            caption,
                        )

            # Delete the placeholder para
            delete_paragraph(p)

        if "{{PARTNERS}}" in p.text:
            # Remove '{{PARTNERS}}'placeholder and add the report content after empty p.text
            p.text = ""

            for category in context["partners"]:
                p.insert_paragraph_before(
                    context["partner_headers"][category], style="Heading14"
                )

                for partner in context["partners"][category]:
                    text, url = partner

                    if url:
                        part = p.insert_paragraph_before(style="HyperlinkList")
                        add_hyperlink(part, url, text)
                    else:
                        p.insert_paragraph_before(style="List Bullet 2").text = text

            # Counties
            if "counties" in context:
                p.insert_paragraph_before("Land Trusts (by county)", style="Heading14")
                for county in context["counties"]:
                    county_name = p.insert_paragraph_before(style="HyperlinkList")
                    add_hyperlink(county_name, county["url"], county["name"])

            # Delete the placeholder para
            delete_paragraph(p)

    report = doc.save(path)

    return report


def generate_report_context(unit_id, config):
    """
        Create json of data points required for report

    Parameters
    ----------
    unit_id: string id of a summary unit
    config: dict of json data from src/config

    Returns
    -------
    context: dict of data in report structure

    """

    with open(os.path.join(DATA_DIR, "{0}.json".format(unit_id))) as json_file:
        data = json.loads(json_file.read())

    total_acres = data["acres"]
    summary_unit_type = "marine lease block" if "M" in unit_id else "subwatershed"

    context = dict()

    context["value"] = {
        "summary_unit_name": data["name"],
        "summary_unit_type": summary_unit_type,
        "acres": "{:,}".format(data["acres"]),
    }
    context["table"] = {}
    context["chart"] = {}
    context["map"] = {}
    context["caption"] = {
        "map_priorities": "Figure 1: Map of Blueprint priority categories within the {0} {1}.".format(
            data["name"], summary_unit_type
        ),
        "chart_priorities": "Figure 2: Proportion of each Blueprint category within the {0} {1}.".format(
            data["name"], summary_unit_type
        ),
        "chart_slr": "Figure 3: Extent of inundation by projected sea level rise within the {0} {1}.".format(
            data["name"], summary_unit_type
        ),
        "chart_urban": "Figure 4: Extent of projected urbanization within the {0} {1}.".format(
            data["name"], summary_unit_type
        )
    }

    # Priorities map

    priorities_config = config["priorities"]
    priorities_map = get_map(unit_id, data, priorities_config)

    context["map"]["priorities"] = priorities_map

    # Priorities table
    table_num = 1

    if data["blueprint"]:
        priorities = dict(col_names=["Priority Category", "Acres", "Percent of Area"])

        priorities["rows"] = [
            (
                priorities_config[str(i)]["label"],
                "{:,}".format(percent_to_acres(percentage, total_acres)),
                "{0}%".format(percentage),
            )
            for i, percentage in enumerate(data.get("blueprint"))
        ]
        priorities["rows"].reverse()

        context["table"]["priorities"] = priorities

        indices = [i for i, v in enumerate(data["blueprint"]) if v > 0]
        indices.reverse()
        values = [data["blueprint"][i] for i in indices]
        colors = [priorities_config[str(i)]["color"] for i in indices]
        labels = [
            "{0} ({1}%)".format(
                priorities_config[str(i)]["label"], data["blueprint"][i]
            )
            for i in indices
        ]
        chart = get_pie_chart(values, colors=colors, labels=labels)
        context["chart"]["priorities"] = chart
        context["caption"]["table_priorities"] = "Table {0}: Extent of each Blueprint priority category within the {1} {2}".format(
            table_num, data["name"], summary_unit_type)
        table_num += 1

    else:
        context["table"]["priorities"] = "No priority information available for this {}".format(summary_unit_type)
        context["chart"]["priorities"] = ""
        context["caption"]["table_priorities"] = ""

    # Ecosystem table

    ecosystems_table = dict(col_names=["Ecosystems", "Acres", "Percent of Area"])
    ecosystems = data.get("ecosystems", [])  # make sure we always have a list
    context["caption"]["table_ecosystems"] = "Table {0}: Extent of each ecosystem within the {1} {2}".format(
        table_num, data["name"], summary_unit_type
        )
    table_num += 1

    ecosystems_with_area = [
        (e, ecosystems[e]["percent"]) for e in ecosystems if "percent" in ecosystems[e]
    ]
    # filter the list to those with percents, and make a new list of tuples of ecosystem ID and percent
    ecosystems_with_area = sorted(
        ecosystems_with_area, key=lambda x: x[1], reverse=True
    )  # sort by percent, from largest percent to smallest

    ecosystems_table["rows"] = [
        (
            config["ecosystems"][e]["label"],
            "{:,}".format(percent_to_acres(percent, total_acres)),
            "{0}%".format(percent),
        )
        for e, percent in ecosystems_with_area
    ]

    context["table"]["ecosystems"] = ecosystems_table

    # Indicators

    context["ecosystem_indicators"] = []

    ind_ecosystems = [
        (key, data["ecosystems"][key].get("percent", None))
        for key in data["ecosystems"]
    ]

    ecosystems_with_percent = [e for e in ind_ecosystems if e[1] is not None]
    ecosystems_with_percent = sorted(
        ecosystems_with_percent, key=lambda x: x[1], reverse=True
    )  # sort descending percent

    ecosystems_without_percent = [e for e in ind_ecosystems if e[1] is None]
    ecosystems_without_percent = sorted(
        ecosystems_without_percent, key=lambda x: x[0]
    )  # sort ascending label

    ecosystems = ecosystems_with_percent + ecosystems_without_percent

    for ecosystem in ecosystems:
        key = ecosystem[0]
        ecosystem_data = data["ecosystems"][key]
        ecosystem_config = config["ecosystems"][key]

        ecosystem_indicators = {
            "ecosystem_name": ecosystem_config["label"],
            "indicators": [],
            "ecosystem_percentage": ecosystem_data.get("percent", ""),
        }

        # Indicator tables

        for indicator in ecosystem_data.get("indicators", []):
            indicator_config = ecosystem_config["indicators"][indicator]
            indicator_data = ecosystem_data["indicators"][indicator]
            indicator_context = {
                "value": {
                    "indicator_name": indicator_config["label"],
                    "indicator_description": indicator_config["description"],
                    "indicator_caption": indicator_config["reportCaption"],
                },
                "table": {
                    "indicator_table": {
                        "col_names": [
                            "",
                            "Indicator Values",
                            "Acres",
                            "Percent of Area",
                        ],
                        "rows": [],
                    }
                },
            }

            good_threshold = indicator_config.get("goodThreshold", None)

            indicator_keys = [int(v) for v in indicator_config["valueLabels"].keys()]
            indicator_keys.sort()  # make sure keys are sorted in ascending order
            indicator_values = list(zip(indicator_keys, indicator_data["percent"]))
            indicator_values.reverse()  # Reverse order so highest priorities at top of table

            if good_threshold is not None:
                # Find threshold position in indicator values so total will be added after
                threshold_position = indicator_keys[::-1].index(good_threshold)

                total_good_percent = 0
                total_not_good_percent = 0
                rows = []

                for label_key, percent in indicator_values:
                    if label_key >= good_threshold:
                        total_good_percent += percent
                        rows.append(
                            [
                                "",
                                indicator_config["valueLabels"][str(label_key)],
                                "{:,}".format(percent_to_acres(percent, total_acres)),
                                "{0}%".format(percent),
                            ]
                        )
                    else:
                        total_not_good_percent += percent
                        rows.append(
                            [
                                "",
                                indicator_config["valueLabels"][str(label_key)],
                                "{:,}".format(percent_to_acres(percent, total_acres)),
                                "{0}%".format(percent),
                            ]
                        )

                # Add lows & highs
                rows[0][0] = "High"
                rows[-1][0] = "Low"

                # Add condition rows

                good_total_row = [
                    "",
                    "Total in good condition",
                    "{:,}".format(percent_to_acres(total_good_percent, total_acres)),
                    "{0}%".format(round(total_good_percent, 1)),
                ]
                not_good_total_row = [
                    "",
                    "Total not in good condition",
                    "{:,}".format(
                        percent_to_acres(total_not_good_percent, total_acres)
                    ),
                    "{0}%".format(round(total_not_good_percent, 1)),
                ]

                # Insert total_good row at threshold and total_not_good at table end
                rows.insert(int(threshold_position) + 1, good_total_row)
                rows.append(not_good_total_row)

                indicator_context["table"]["indicator_table"]["rows"] = rows

            else:
                rows = [
                    [
                        "",
                        indicator_config["valueLabels"][str(label_key)],
                        "{:,}".format(percent_to_acres(percent, total_acres)),
                        "{0}%".format(percent),
                    ]
                    for label_key, percent in indicator_values
                ]

                # Add lows & highs
                rows[0][0] = "High"
                rows[-1][0] = "Low"

                indicator_context["table"]["indicator_table"]["rows"] = rows

            ecosystem_indicators["indicators"].append(indicator_context)

        context["ecosystem_indicators"].append(ecosystem_indicators)

    # Partners list - name and url separated by categories

    context["partner_headers"] = {
        "regional": "Regional Conservation Plans",
        "state": "Statewide Conservation Plans",
        "marine": "Marine Conservation Plans",
    }

    context["partners"] = defaultdict(list)
    for plan_key in data["plans"]:
        plan = config["plans"][plan_key]
        label = plan["label"]
        url = plan.get("url", "")
        context["partners"][plan["type"]].append((label, url))

    # Counties list - name and url

    if "counties" in data:
        context["counties"] = [
            {"name": value, "url": "http://findalandtrust.org/counties/{0}".format(key)}
            for key, value in data["counties"].items()
        ]  # key is FIPS and value is county with state

    # Ownership table

    if "owner" in data:
        owners = dict(col_names=["Ownership", "Acres", "Percent of Area"])

        owners["rows"] = [
            [
                config["owners"][key]["label"],
                "{:,}".format(percent_to_acres(percent, total_acres)),
                "{0}%".format(percent),
            ]
            for key, percent in data["owner"].items()
        ]

        own_perc_sum = sum(data["owner"].values())

        if own_perc_sum < 100:
            perc_remainder = round(100 - own_perc_sum, 1)
            owners["rows"].append(
                [
                    "Not conserved",
                    "{:,}".format(percent_to_acres(perc_remainder, total_acres)),
                    "{0}%".format(perc_remainder),
                ]
            )

        context["table"]["ownership"] = owners
    else:
        context["table"]["ownership"] = "No information available"

    # Protections table

    protection = dict(col_names=["Land Protection Status", "Acres", "Percent of Area"])

    if "gap" in data:
        protection["rows"] = [
            [
                config["protection"][key]["label"],
                "{:,}".format(percent_to_acres(percent, total_acres)),
                "{0}%".format(percent),
            ]
            for key, percent in data["gap"].items()
        ]

        pro_perc_sum = sum(data["gap"].values())

        if pro_perc_sum < 100:
            perc_remainder = round(100 - pro_perc_sum, 1)
            protection["rows"].append(
                [
                    "Not conserved",
                    "{:,}".format(percent_to_acres(perc_remainder, total_acres)),
                    "{0}%".format(perc_remainder),
                ]
            )

        context["table"]["protection"] = protection
    else:
        context["table"]["protection"] = "No information available"

    # Threats

    if "slr" in data and data["slr"]:
        chart = get_line_chart(
            config["slr"],
            data["slr"],
            x_label="Amount of sea level rise (feet)",
            y_label="Percent of area",
            color="#004da8",
            alpha=0.3,
        )
        context["chart"]["slr"] = chart
    else:
        context["chart"]["slr"] = "No sea level rise data available"

    if "urban" in data and data["urban"]:
        chart = get_line_chart(
            config["urbanization"],
            data["urban"],
            x_label="Decade",
            y_label="Percent of area",
            color="#D90000",
            alpha=0.3,
        )
        context["chart"]["urban"] = chart
    else:
        context["chart"]["urban"] = "No urban growth data available"

    # Table captions for last two tables, Ownership and Protection status

    for ecosystem in context["ecosystem_indicators"]:
        for indicator in ecosystem.get("indicators", []):
            table_num += 1

    if "rows" in context["table"]["ownership"]:
        context["caption"]["table_ownership"] = "Table {0}: Extent of ownership class within the {1} {2}.".format(
            table_num, data["name"], summary_unit_type
        )
        table_num += 1
    else:
        context["caption"]["table_ownership"] = ""

    if "rows" in context["table"]["protection"]:
        context["caption"]["table_protection"] = "Table {0}: Extent of land protection status within the {1} {2}.".format(
            table_num, data["name"], summary_unit_type
        )
        table_num += 1
    else:
        context["caption"]["table_protection"] = ""

    return context


def percent_to_acres(percentage, total_acres):
    return round(total_acres * (percentage / 100))


def create_table(doc, data, para):
    """
    Builds table at end of doc

    Parameters
    ----------
    doc: docx.Document object
    data: dict
        table data; 'data' contains a list of column names ('col_names') and a list of rows ('rows'),
        where each row is a list of values for each column
    para: paragraph object
        Where the table needs to be moved

    Returns
    -------
    end_buffer
        location that the last item was added to the document

    """
    # TODO: ensure tables don't run over a page break

    styles = doc.styles
    ncols = len(data["col_names"])

    # Create table with one row for column headings
    table = doc.add_table(rows=1, cols=ncols)
    table.style = styles[DOC_STYLE]

    for index, heading in enumerate(data["col_names"]):
        cell = table.cell(0, index)
        cell.text = heading
        # Change justification of first heading
        if ncols is 3 and index is 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif ncols is 4 and index is 1:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for datarow in data["rows"]:
        row = table.add_row()
        condition_style = (
            str(datarow[1]) == "Total in good condition"
            or str(datarow[1]) == "Total not in good condition"
        )
        for index, datacell in enumerate(datarow):
            row.cells[index].text = str(datacell)
            if ncols is 4 and index is 1:
                row.cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            if condition_style is True:
                # A new shade must be generated for each cell
                new_shade = shade_generator(TOTAL_ROW_COLOR)
                row.cells[index]._tc.get_or_add_tcPr().append(new_shade)

    if "Indicator Values" in data["col_names"]:
        table_type = "indicator"
    else:
        table_type = "other"

    set_col_widths(table, table_type)

    _move_table_after(table, para)

    end_buffer = doc.add_paragraph("")
    _move_p_after_t(table, end_buffer)

    return end_buffer


def shade_generator(color):
    return docx.oxml.parse_xml(
        r'<w:shd {0} w:fill="{1}"/>'.format(docx.oxml.ns.nsdecls("w"), color)
    )


def set_col_widths(table, table_type):
    """
    Set widths for standard table columns

    Parameters
    ----------
    table: table object

    """
    if table_type is "indicator":
        widths = (Cm(1), Cm(9), Cm(3.5), Cm(3.5))
    else:
        widths = (Cm(9), Cm(4), Cm(4))

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def _move_p_after(p_move, p_destination):
    """
    Move one paragraph after another

    Parameters
    ----------
    heading
        heading to be moved
    paragraph
        location where table will be moved

    """
    p_destination._p.addnext(p_move._p)


def _move_table_after(table, paragraph):
    """
    Move the table after 'paragraph.'
    Normally, tables are created at the end of the document.
    This function allows you to move them to a different location in the document.
    The paragraph passed in is not modified, it is just used as a reference point for inserting the table.

    Parameters
    ----------
    table
        table to be moved
    paragraph
        location where table will be moved

    """
    paragraph._p.addnext(table._tbl)


def _move_p_after_t(table, para):
    """
    Move a paragraph after a table

    Parameters
    ----------
    table
        table to be moved
    para: paragraph object
        location where table will be moved

    """
    table._tbl.addnext(para._p)


def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    Parameters
    ----------
    paragraph: The paragraph we are adding the hyperlink to.
    url: A string containing the required url
    text: The text displayed for the url

    Returns
    -------
    hyperlink: hyperlink object

    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement("w:r")

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def get_map(unit_id, data, priorities):
    bounds = data["bounds"]

    # convert priorities to legend, in descending priority (dropping Not a Priority class)
    legend = [
        (int(k), v["color"], v["label"]) for k, v in priorities.items() if k != "0"
    ]
    legend = sorted(legend, key=lambda x: x[0], reverse=True)
    # strip off the value used for sorting
    legend = [entry[1:] for entry in legend]

    m = get_map_image(unit_id, bounds, legend)

    buffer = BytesIO()
    m.save(buffer, "PNG")
    return buffer


def _resolve(scope, key, context):
    """
    Resolve scope and key to a context item
    Provides very minimal validation of presence

    Parameters
    ----------
    scope: str
        singular variant of scope in context
    key: str
        key to lookup context item in context within scope
    context: dict

    Returns
    -------
    context item: str, dict, etc
    """

    if scope not in context:
        raise ValueError("Scope {0} is not found in context!".format(scope))

    if key not in context[scope]:
        raise ValueError("Key {0} is not found in context!".format(key))

    return context[scope][key]


if __name__ == "__main__":
    id = "I1"
    outpath = "/tmp/{0}.docx".format(id)
    create_report(id, outpath)
